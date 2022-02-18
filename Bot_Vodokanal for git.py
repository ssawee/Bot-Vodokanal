#Основная цель данного бота - создавать таблицу и записывать в нее полученные от пользователей данные 
#Подключение необходимых модулей
import re
import sys
import random
import datetime
import docx
import vk_api
from docx.enum.text import WD_ALIGN_PARAGRAPH
from vk_api.longpoll import VkLongPoll, VkEventType
from vk_api.bot_longpoll import VkBotLongPoll, VkBotEventType

#Ввод актуального токена и id группы для авторизации (желательно подключить файл с данными)
token=token
group=group


#Авторизация "Вконтакте" как сообщество
vk_session=vk_api.VkApi(token=token)
print('vk_session')
vk_session._auth_token()
print('auth_token()')
#Работа с сообщениями
longpoll=VkBotLongPoll(vk_session, group)
print('longpoll')
vk=vk_session.get_api()
print('get_api()')


#Функция создания таблицы
def generate_table():

    #Генерация таблицы
    document=docx.Document()

    #Параметры титульного текста
    paragraph=document.add_paragraph('')
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    title=paragraph.add_run('Итоговая таблица записи показаний')
    title.bold=True
    title.font.size=304800

    #Параметры таблицы
    rows=1
    cols=4
    table=document.add_table(rows = rows, cols = cols)
    table.style='Table Grid'

    list=['Номер', 'Имя (id)', 'Дата/Время', 'Показатели']

    #Построение таблицы
    for row in range(1):
        for col in range(int(cols)):
            cell=table.cell(row, col)
            cell.text=list[col]

    document.save('table.docx')


#Функция заполнения таблицы    
def completion_table(data):

    #Заполнение таблицы
    document=docx.Document('table.docx')

    document.tables[0]

    cols=4
    
    for row in range(int(counter), int(counter+1)):
        document.tables[0].add_row()
        for col in range(cols):
            document.tables[0].cell(row, col).text = data[col]

    document.save('table.docx')


#Функция получения цифровых данных для таблицы 
def numbers_from_string(response):

    #Получение цифровых данных из сообщения
    indications=''.join(re.findall(r'\d+', response)) 

    return indications


#Функция получения текстовых данных для таблицы 
def get_name(from_id):

    #Получение имени и фамилии пользователя
    user=vk_session.method("users.get", {"user_ids": from_id})
    fullname=user[0]['first_name']+' '+user[0]['last_name']

    return fullname


#Функция преобразование JSON-пакета в необходимый массив данных
def handler(objects, counter):

    indications=numbers_from_string(response)
    from_id=event.object.from_id
    get_name(from_id)
    time=str(datetime.datetime.date(datetime.datetime.now()).strftime("%d.%m.%y"))+'/'+str(datetime.datetime.time(datetime.datetime.now()).strftime("%H:%M:%S"))

    data=[str(counter), str(get_name(from_id)), str(time), str(indications)]

    return data


#Класс независимого счётчика
class IncrementCounter:
    
    def __init__(self):
        self._value=0
    
    def new_value(self):
        self._value+=1
        return self._value


counter1=IncrementCounter()

generate_table()


#Основной цикл
while True:
    for event in longpoll.listen():
        if event.type==VkBotEventType.MESSAGE_NEW:
            #Работа с чатами
            if event.from_chat:
                print('\nНовое сообщение')
                print(datetime.datetime.date(datetime.datetime.now()).strftime("%d.%m.%y"))
                print(datetime.datetime.time(datetime.datetime.now()).strftime("%H:%M:%S"))
                print('chat_id:',event.chat_id)
                print('user_id:',event.object.from_id)
                print('message:',event.object.text)
                print('\n',event.object)
                chat_id=event.chat_id
                random_id=vk_api.utils.get_random_id()
                #Работа с текстом
                if event.object.text:
                    response=event.object.text.lower()
                    if response=='/справка':
                        vk.messages.send(chat_id=chat_id, message='Команда "/показания" используется в формате: \n/показания 12345', random_id=random_id)
                    if '/показания' in response.split(' '):
                        if numbers_from_string(response)!='':
                            objects=event.object
                            counter=counter1.new_value()
                            handler(objects, counter)
                            completion_table(handler(objects, counter))
                            vk.messages.send(chat_id=chat_id, message='Ваши показания записаны!', random_id=random_id)
                        else:
                            vk.messages.send(chat_id=chat_id, message='Ошибка ввода данных! Проверьте корректность вводимых показаний!', random_id=random_id)
